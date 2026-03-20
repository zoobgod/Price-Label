"""Build output documents: Price List and Label(s).

Template filling now works at the *run* level so that the original
formatting (fonts, sizes, bold, colours) is preserved.
"""

from __future__ import annotations

import re
from collections import OrderedDict
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from .models import ExtractedData, Position

DEFAULT_STORAGE_TEMPERATURE_EN = "+15C to +25C ambient"
DEFAULT_STORAGE_TEMPERATURE_RU = "+15C до +25C"

# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def _price_str(value: float | None) -> str:
    if value is None:
        return ""
    return f"{value:,.2f}"


def _qty_str(value: float | None) -> str:
    if value is None:
        return ""
    return str(int(value)) if float(value).is_integer() else f"{value:g}"


def _money(value: float | None, currency: str) -> str:
    s = _price_str(value)
    return f"{s} {currency}".strip() if s else ""


def _normalize_temp(temp: str) -> str:
    t = _clean(temp)
    if not t:
        return ""
    m = re.search(
        r"([+-]?\d{1,2})\s*°?\s*C\s*(?:to|до|–|-|~)\s*([+-]?\d{1,2})\s*°?\s*C",
        t,
        re.IGNORECASE,
    )
    if m:
        lo, hi = int(m.group(1)), int(m.group(2))
        fmt = lambda v: f"+{v}" if v >= 0 else str(v)
        suffix = " ambient" if "ambient" in t.lower() else ""
        return f"{fmt(lo)}C to {fmt(hi)}C{suffix}"
    if "room temperature" in t.lower() or "ambient" in t.lower():
        return DEFAULT_STORAGE_TEMPERATURE_EN
    return t


def _temp_ru(temp_en: str) -> str:
    t = _normalize_temp(temp_en) or DEFAULT_STORAGE_TEMPERATURE_EN
    t = t.replace(" to ", " до ").replace(" ambient", "")
    return t


def _format_date(value: str) -> str:
    raw = _clean(value)
    if not raw:
        return ""
    for fmt in (
        "%d-%b-%y", "%d-%b-%Y", "%d.%m.%Y", "%d.%m.%y",
        "%d/%m/%Y", "%d/%m/%y", "%d-%m-%Y", "%d-%m-%y",
    ):
        for candidate in (raw, raw.replace("/", "-").replace(".", "-")):
            try:
                return datetime.strptime(candidate, fmt).strftime("%d.%m.%y")
            except ValueError:
                pass
    m = re.search(r"(\d{2})[./-](\d{2})[./-](\d{2,4})", raw)
    if m:
        yr = m.group(3)[-2:]
        return f"{m.group(1)}.{m.group(2)}.{yr}"
    return raw


def _normalize_terms(value: str) -> str:
    text = _clean(value)
    if not text:
        return ""
    m = re.search(r"\b(CPT|FOB|CIF|EXW|DAP|DDP|FCA)\b", text, re.IGNORECASE)
    if not m:
        return text
    incoterm = m.group(1).upper()
    low = text.lower()
    parts = [incoterm]
    if re.search(r"\bby\s+air\b", low):
        parts.append("BY AIR")
    # Extract city – take first capitalised word after incoterm that isn't a keyword
    city_match = re.search(
        r"\b(MOSCOW|HYDERABAD|DELHI|MUMBAI|CHENNAI|KOLKATA|BANGALORE|DUBAI|ISTANBUL|LONDON|BERLIN|PARIS|TOKYO)\b",
        text,
        re.IGNORECASE,
    )
    if city_match:
        parts.append(city_match.group(1).upper())
    return " ".join(parts)


# ---------------------------------------------------------------------------
# Run-level template filling (preserves formatting)
# ---------------------------------------------------------------------------

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}")


def _replace_in_paragraph(paragraph, context: dict[str, str]) -> int:
    """Replace {{KEY}} placeholders while preserving run formatting.

    Strategy: concatenate all run texts, do replacements on the joined
    string, then redistribute text back across runs (first run gets the
    new text, remaining runs are emptied). This preserves the *first*
    run's formatting for the whole paragraph – acceptable because
    within a template paragraph the formatting is almost always uniform.
    """
    runs = paragraph.runs
    if not runs:
        return 0

    full = "".join(r.text or "" for r in runs)
    if "{{" not in full:
        return 0

    replaced = 0
    new = full
    for key, val in context.items():
        placeholder = "{{" + key + "}}"
        if placeholder in new:
            new = new.replace(placeholder, val)
            replaced += 1

    # Wipe unresolved placeholders
    new = _PLACEHOLDER_RE.sub("", new)
    # Clean stray artefacts
    new = re.sub(r"\(\s*,\s*\)", "()", new)
    new = re.sub(r"\s{2,}", " ", new).strip()

    if new == full:
        return 0

    runs[0].text = new
    for r in runs[1:]:
        r.text = ""
    return replaced


def _fill_template(template_bytes: bytes, context: dict[str, str]) -> tuple[bytes, int]:
    """Fill a .docx template, preserving formatting. Returns (bytes, count_of_replaced)."""
    doc = Document(BytesIO(template_bytes))
    total = 0

    for para in doc.paragraphs:
        total += _replace_in_paragraph(para, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += _replace_in_paragraph(para, context)

    # Also handle headers / footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is None:
                continue
            for para in header_footer.paragraphs:
                total += _replace_in_paragraph(para, context)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue(), total


# ---------------------------------------------------------------------------
# Context builder
# ---------------------------------------------------------------------------


def _build_context(
    data: ExtractedData,
    positions: list[Position],
    company_profile: dict[str, str] | None,
    temperature_en: str,
) -> dict[str, str]:
    cp = company_profile or {}

    exporter_en = cp.get("exporter_company_name_en") or data.exporter_name or ""
    exporter_ru = cp.get("exporter_company_name_ru") or data.exporter_name_ru or exporter_en
    exporter_addr = cp.get("exporter_company_address_en") or data.exporter_address or ""

    storage_en = _normalize_temp(temperature_en) or DEFAULT_STORAGE_TEMPERATURE_EN
    storage_ru = cp.get("storage_temperature_ru") or _temp_ru(storage_en)

    ctx: dict[str, str] = {
        "INVOICE_NO": data.invoice_no,
        "INVOICE_DATE": _format_date(data.invoice_date),
        "TERMS_OF_DELIVERY": _normalize_terms(data.terms_of_delivery),
        "PERIOD_OF_VALIDITY": data.period_of_validity or "",
        "SPECIFICATION_DATE": _format_date(data.specification_date),
        "STORAGE_TEMPERATURE": storage_en,
        "STORAGE_TEMPERATURE_EN": storage_en,
        "STORAGE_TEMPERATURE_RU": storage_ru,
        "EXPORTER_COMPANY_NAME_EN": exporter_en,
        "EXPORTER_COMPANY_NAME_RU": exporter_ru,
        "EXPORTER_COMPANY_ADDRESS_EN": exporter_addr,
        "EXPORTER_COMPANY_ADRESS_EN": exporter_addr,  # typo-compat with user template
    }

    default_currency = data.currency or ""

    for i, pos in enumerate(positions, start=1):
        qty = _qty_str(pos.quantity)
        currency = pos.currency or default_currency
        ctx[f"POSITION_{i}_NAME_EN"] = pos.name_en or ""
        ctx[f"POSITION_{i}_NAME_RU"] = pos.name_ru or pos.name_en or ""
        ctx[f"POSITION_{i}_QUANTITY"] = qty
        ctx[f"POSITION_{i}_QUANTITY_EN"] = f"{qty} un" if qty else ""
        ctx[f"POSITION_{i}_QUANTITY_RU"] = f"{qty} шт" if qty else ""
        ctx[f"POSITION_{i}_PACKING"] = pos.packing_en or ""
        ctx[f"POSITION_{i}_PACKING_EN"] = pos.packing_en or ""
        ctx[f"POSITION_{i}_PACKING_RU"] = pos.packing_ru or pos.packing_en or ""
        ctx[f"POSITION_{i}_UNIT_PRICE"] = _money(pos.unit_price, currency)
        ctx[f"POSITION_{i}_TOTAL_PRICE"] = _money(pos.total_price, currency)
        ctx[f"POSITION_{i}_CURRENCY"] = currency

    return ctx


# ---------------------------------------------------------------------------
# Default document generators (used when no template is uploaded)
# ---------------------------------------------------------------------------


def _set_font(doc: Document, name: str = "Times New Roman", size: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)


def _default_price_doc(data: ExtractedData, positions: list[Position], cp: dict[str, str] | None) -> bytes:
    doc = Document()
    _set_font(doc)

    p = doc.add_paragraph("On the company letterhead")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph("PRICE LIST")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    dp = doc.add_paragraph(_format_date(data.invoice_date))
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for idx, pos in enumerate(positions):
        if idx > 0:
            doc.add_paragraph()
        currency = pos.currency or data.currency or ""
        doc.add_paragraph(f"PRODUCT NAME: {pos.name_en or '-'}")
        doc.add_paragraph(f"QUANTITY: {_qty_str(pos.quantity) or '-'}")
        doc.add_paragraph(f"PACKING: {pos.packing_en or '-'}")
        doc.add_paragraph(f"UNIT PRICE: {_money(pos.unit_price, currency) or '-'}")
        doc.add_paragraph(f"TOTAL AMOUNT: {_money(pos.total_price, currency) or '-'}")

    doc.add_paragraph()
    doc.add_paragraph(f"TERMS OF DELIVERY: {_normalize_terms(data.terms_of_delivery) or '-'}")
    doc.add_paragraph(f"PERIOD OF VALIDITY: {data.period_of_validity or '-'}")
    doc.add_paragraph()
    doc.add_paragraph("stamp / signature")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _default_label_doc(
    data: ExtractedData,
    positions: list[Position],
    temperature_en: str,
    cp: dict[str, str] | None,
) -> bytes:
    ctx = _build_context(data, positions, cp, temperature_en)
    doc = Document()
    _set_font(doc)

    hdr = doc.add_paragraph("Наименование товара / Product name - Quantity / Кол-во:")
    hdr.runs[0].bold = True

    for i, _ in enumerate(positions, start=1):
        name_en = ctx.get(f"POSITION_{i}_NAME_EN", "")
        name_ru = ctx.get(f"POSITION_{i}_NAME_RU", "")
        qty_en = ctx.get(f"POSITION_{i}_QUANTITY_EN", "")
        qty_ru = ctx.get(f"POSITION_{i}_QUANTITY_RU", "")
        pack_en = ctx.get(f"POSITION_{i}_PACKING_EN", "")
        pack_ru = ctx.get(f"POSITION_{i}_PACKING_RU", "")
        line = f"{name_en} / {name_ru} {qty_en} ({pack_en}) / {qty_ru} ({pack_ru})"
        doc.add_paragraph(_clean(line))

    doc.add_paragraph(
        f"Shipping Conditions: Require temperature-controlled shipping "
        f"must be kept between ({ctx['STORAGE_TEMPERATURE_EN']})"
    )
    doc.add_paragraph(
        f"Условия транспортировки: требуется контролируемая температура "
        f"при транспортировке ({ctx['STORAGE_TEMPERATURE_RU']})"
    )
    doc.add_paragraph(
        f'Shipper/Отправитель: "{ctx["EXPORTER_COMPANY_NAME_EN"]}" '
        f'/ «{ctx["EXPORTER_COMPANY_NAME_RU"]}»'
    )
    doc.add_paragraph(
        f"Contact information / Контактная информация: {ctx['EXPORTER_COMPANY_ADRESS_EN']}"
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def generate_price_list_doc(
    data: ExtractedData,
    template_bytes: bytes | None = None,
    company_profile: dict[str, str] | None = None,
) -> bytes:
    positions = data.positions or [Position()]
    temp = positions[0].storage_temperature or data.storage_temperature or DEFAULT_STORAGE_TEMPERATURE_EN
    ctx = _build_context(data, positions, company_profile, temp)

    if template_bytes:
        filled, count = _fill_template(template_bytes, ctx)
        if count > 0:
            return filled

    return _default_price_doc(data, positions, company_profile)


def _group_by_temperature(data: ExtractedData) -> OrderedDict[str, list[Position]]:
    grouped: OrderedDict[str, list[Position]] = OrderedDict()
    for pos in data.positions or [Position()]:
        temp = _normalize_temp(pos.storage_temperature) or _normalize_temp(data.storage_temperature) or DEFAULT_STORAGE_TEMPERATURE_EN
        grouped.setdefault(temp, []).append(pos)
    return grouped


def _temp_slug(temp: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", temp).strip("_")[:60] or "ambient"


def generate_label_docs_by_temperature(
    data: ExtractedData,
    template_bytes: bytes | None = None,
    company_profile: dict[str, str] | None = None,
) -> OrderedDict[str, bytes]:
    grouped = _group_by_temperature(data)
    docs: OrderedDict[str, bytes] = OrderedDict()
    many = len(grouped) > 1

    for temperature, positions in grouped.items():
        ctx = _build_context(data, positions, company_profile, temperature)
        label_bytes: bytes | None = None

        if template_bytes:
            filled, count = _fill_template(template_bytes, ctx)
            if count > 0:
                label_bytes = filled

        if label_bytes is None:
            label_bytes = _default_label_doc(data, positions, temperature, company_profile)

        filename = f"Label_{_temp_slug(temperature)}.docx" if many else "Label.docx"
        docs[filename] = label_bytes

    return docs
